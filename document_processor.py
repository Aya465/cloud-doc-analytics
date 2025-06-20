import os
import fitz  # PyMuPDF
import docx
import re

def extract_title_from_pdf(filepath):
    doc = fitz.open(filepath)
    for page in doc:
        text = page.get_text().strip()
        if text:
            return text.split('\n')[0]  # First line as title
    return os.path.basename(filepath)

def extract_title_from_docx(filepath):
    doc = docx.Document(filepath)
    if doc.paragraphs:
        return doc.paragraphs[0].text.strip()
    return os.path.basename(filepath)

def extract_text_from_document(filepath):
    if filepath.endswith(".pdf"):
        doc = fitz.open(filepath)
        return " ".join([page.get_text() for page in doc])
    elif filepath.endswith(".docx"):
        doc = docx.Document(filepath)
        return " ".join([para.text for para in doc.paragraphs])
    return ""

def search_in_document(filepath, keyword):
    results = []
    if filepath.endswith(".pdf"):
        doc = fitz.open(filepath)
        for page_num, page in enumerate(doc):
            text_instances = page.search_for(keyword)
            if text_instances:
                results.append((page_num + 1, text_instances))
    elif filepath.endswith(".docx"):
        doc = docx.Document(filepath)
        for i, para in enumerate(doc.paragraphs):
            if keyword.lower() in para.text.lower():
                results.append((i + 1, para.text))
    return results

def sort_documents_by_title(filepaths):
    titles = []
    for path in filepaths:
        if path.endswith(".pdf"):
            title = extract_title_from_pdf(path)
        elif path.endswith(".docx"):
            title = extract_title_from_docx(path)
        else:
            continue
        titles.append((title, path))
    return [x[1] for x in sorted(titles, key=lambda x: x[0].lower())]
